#!/usr/bin/env python3
"""
Post-process a generated DOCX to:
- set all tables to 100% width
- reduce table font size slightly (e.g., 11 pt)

Usage:
  python3 scripts/conversion/postprocess_docx_tables.py /path/to/file.docx

This script uses the vendored python-docx in tools/vendor if available.
"""
from pathlib import Path
import sys

# Prefer vendored python-docx if present
VENDOR = Path(__file__).resolve().parents[1] / 'tools' / 'vendor'
if VENDOR.exists():
    sys.path.insert(0, str(VENDOR))

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception as e:
    print("[postprocess_docx_tables] python-docx not available:", e)
    sys.exit(0)  # do not fail the build


def set_table_width_pct(tbl, pct: int = 100):
    """Set a table width to a percentage (default 100%).
    Uses low-level oxml to be compatible with older python-docx versions.
    Word uses fiftieths of a percent for w attribute when type='pct'.
    """
    # Ensure <w:tblPr>
    tblPr = getattr(tbl._tbl, 'tblPr', None)
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        # Insert as first child of <w:tbl>
        tbl._tbl.insert(0, tblPr)

    # Ensure/Update <w:tblW>
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    # 100% -> 5000 (50 * percent)
    tblW.set(qn('w:w'), str(int(pct * 50)))


def reduce_table_font(tbl, pt_size: float = 11.0):
    """Set a slightly smaller font size for all runs inside a table."""
    size = Pt(pt_size)
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    try:
                        run.font.size = size
                    except Exception:
                        pass


def process(docx_path: Path):
    doc = Document(str(docx_path))
    changed = False
    for t in doc.tables:
        set_table_width_pct(t, 100)
        reduce_table_font(t, 11.0)
        changed = True
    if changed:
        doc.save(str(docx_path))
        print(f"[postprocess_docx_tables] updated tables in: {docx_path}")
    else:
        print(f"[postprocess_docx_tables] no tables found: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: postprocess_docx_tables.py <docx-file>")
        sys.exit(1)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"DOCX not found: {path}")
        sys.exit(1)
    process(path)
