#!/usr/bin/env python3
"""
Generate the reference DOCX used by pandoc exports.

The resulting file lives at templates/reference.docx and contains:
* Times New Roman as the base font for body text.
* Black headings with Times New Roman applied.
* Clearly delimited sample sections so authors can tweak styles quickly.

Run this script whenever you need to refresh the reference document.
"""

from pathlib import Path
import sys
from typing import Dict

REPO_ROOT = Path(__file__).resolve().parent.parent
VENDOR_PATH = REPO_ROOT / "tools" / "vendor"
CONFIG_PATH = REPO_ROOT / "config" / "style_overrides.txt"

if VENDOR_PATH.exists():
    sys.path.insert(0, str(VENDOR_PATH))

from docx import Document  # type: ignore  # pylint: disable=import-error
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt, RGBColor  # type: ignore


def load_overrides(config_path: Path) -> Dict[str, str]:
    """Read key=value overrides from a text file, ignoring comments."""
    if not config_path.exists():
        return {}

    overrides: Dict[str, str] = {}
    for raw_line in config_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        overrides[key.strip().lower()] = value.strip()
    return overrides


def set_style_font(style, font_name, size_pt=None, color_rgb=None, bold=None):
    """Configure a style font with optional size, color, and weight."""
    font = style.font
    font.name = font_name
    if size_pt is not None:
        font.size = Pt(size_pt)
    if color_rgb is not None:
        font.color.rgb = RGBColor(*color_rgb)
    if bold is not None:
        font.bold = bold

    # Ensure east Asian and complex script fonts follow the same family.
    rpr = style.element.get_or_add_rPr()
    for key in ("w:eastAsia", "w:cs", "w:hAnsi"):
        rpr.rFonts.set(qn(key), font_name)


def main() -> None:
    doc = Document()

    overrides = load_overrides(CONFIG_PATH)

    font_family = overrides.get("font.family", "Times New Roman")
    justify_body = overrides.get("body.justify", "true").lower() in {"true", "1", "yes", "on"}

    default_title_spacing = 18.0
    try:
        title_spacing_after = float(overrides.get("title.spacing_after_pt", default_title_spacing))
    except ValueError:
        title_spacing_after = default_title_spacing

    styles = doc.styles

    # Normal/body style
    normal_style = styles["Normal"]
    set_style_font(normal_style, font_family, size_pt=12, color_rgb=(0, 0, 0))
    normal_style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify_body else WD_ALIGN_PARAGRAPH.LEFT
    )

    # Heading styles: keep the default hierarchy but force consistent font + black text.
    for level, size in [(1, 18), (2, 16), (3, 14)]:
        heading_style = styles[f"Heading {level}"]
        set_style_font(heading_style, font_family, size_pt=size, color_rgb=(0, 0, 0), bold=True)

    # Title/Subtitle for the intro section, matched with Times New Roman.
    if "Title" in styles:
        title_style = styles["Title"]
        set_style_font(title_style, font_family, size_pt=22, color_rgb=(0, 0, 0), bold=True)
        title_style.paragraph_format.space_after = Pt(title_spacing_after)
    if "Subtitle" in styles:
        subtitle_style = styles["Subtitle"]
        set_style_font(subtitle_style, font_family, size_pt=14, color_rgb=(0, 0, 0))

    doc.add_paragraph("Playground font e stili", style="Title")
    doc.add_paragraph(
        "Modifica i blocchi qui sotto per personalizzare gli stili della reference DOCX usata da pandoc.",
        style="Subtitle",
    )

    doc.add_paragraph("=== BLOCCO NORMAL: INIZIO ===", style="Normal")
    doc.add_paragraph(
        "Questo paragrafo usa lo stile 'Normal'. Cambia font, dimensione, colore e spaziatura del corpo qui.",
        style="Normal",
    )
    doc.add_paragraph("=== BLOCCO NORMAL: FINE ===", style="Normal")

    doc.add_paragraph("=== BLOCCO HEADING 1 ===", style="Heading 1")
    doc.add_paragraph(
        "Il testo immediatamente sopra usa 'Heading 1'. Imposta qui colore (nero), font e altri attributi.",
        style="Normal",
    )

    doc.add_paragraph("=== BLOCCO HEADING 2 ===", style="Heading 2")
    doc.add_paragraph(
        "Il titolo sopra usa 'Heading 2'. Puoi modificare questo blocco per regolare stili di secondo livello.",
        style="Normal",
    )

    doc.add_paragraph("=== BLOCCO HEADING 3 ===", style="Heading 3")
    doc.add_paragraph(
        "Anche gli heading di terzo livello sono preimpostati su Times New Roman e colore nero.",
        style="Normal",
    )

    doc.add_paragraph(
        "Suggerimento: modifica i testi delimitati per vedere l'effetto delle tue personalizzazioni quando rigeneri gli output.",
        style="Normal",
    )

    output_path = REPO_ROOT / "templates" / "reference.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ Reference DOCX aggiornato: {output_path}")


if __name__ == "__main__":
    main()
